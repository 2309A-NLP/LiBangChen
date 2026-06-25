# -*- coding: utf-8 -*-
"""
文件上传与解析服务
功能：处理用户上传文件的持久化、解析、分块和向量化同步。
支持多种文件格式：DOCX、XLSX、CSV、JSON、TXT、MD、图片（OCR）。
当 Milvus 可用时，自动将文件分块同步到向量库。

主要类：UserFileService
  - upload_file(): 上传并解析文件
  - analyze_file_role(): 分析文件内容推荐角色
  - list_conversation_files(): 列出会话文件
  - delete_uploaded_file(): 删除文件
  - recommend_role(): 基于关键词匹配推荐角色
"""

import csv  # CSV 文件解析
import base64  # Base64 编码，用于图片转码
import io  # 内存文件操作
import json  # JSON 解析
import os  # 操作系统接口
import re  # 正则表达式
import subprocess  # 子进程管理，用于运行 OCR worker
import sys  # 系统参数
import uuid  # 生成唯一文件名
from pathlib import Path  # 路径操作
from typing import Dict, List  # 类型注解

import docx  # DOCX 文件解析
import fitz
import openpyxl  # XLSX 文件解析
import pdfplumber
from PIL import Image

from config import UPLOAD_CONFIG  # 上传配置
from data_processor import DataProcessor  # 数据处理器
from llm_settings import build_multimodal_openai_client, load_multimodal_llm_config  # 在线多模态客户端
from logging_utils import get_logger  # 日志工具
from models import Conversation, UploadedFile, UserDocumentChunk  # 数据库模型

logger = get_logger(__name__)  # 获取当前模块的日志记录器


class UserFileService:
    """
    文件上传与解析服务类。
    
    处理用户上传文件的持久化、解析、分块和向量化同步。
    支持多种文件格式：DOCX、XLSX、CSV、JSON、TXT、MD、图片（OCR）。
    当 Milvus 可用时，自动将文件分块同步到向量库。
    """

    # 各角色对应的关键词列表，用于文件内容匹配推荐角色
    ROLE_KEYWORDS = {
        "lawyer": [  # 律师角色关键词
            "合同", "劳动", "仲裁", "诉讼", "证据", "违约", "赔偿", "法院", "法律", "用人单位",
        ],
        "stock_analyst": [  # 股票分析师角色关键词
            "财报", "利润", "营收", "现金流", "应收账款", "存货", "估值", "股票", "投资", "负债",
        ],
        "teacher": [  # 教师角色关键词
            "学生", "考试", "复习", "作业", "课程", "数学", "英语", "知识点", "学习", "教学",
        ],
        "psychological_counselor": [  # 心理咨询师角色关键词
            "焦虑", "抑郁", "情绪", "失眠", "压力", "心理", "咨询", "自伤", "自杀", "睡眠",
        ],
        "doctor": [  # 医生角色关键词
            "发热", "咳嗽", "血压", "血糖", "症状", "医院", "就医", "药物", "诊断", "体检",
        ],
        "scientist": [  # 科学家角色关键词
            "实验", "假设", "变量", "对照", "数据", "论文", "研究", "样本", "结论", "可重复",
        ],
    }

    # 多模态解析结果中需要过滤的无效 token
    INVALID_MULTIMODAL_TOKENS = {
        "imstart", "im_end", "imend", "<|im_start|>", "<|im_end|>",
    }

    def __init__(self, db_session):
        """初始化文件上传服务。"""
        self.db = db_session  # 数据库会话
        self.processor = DataProcessor()  # 数据处理器实例
        self.upload_root = Path(UPLOAD_CONFIG["root_dir"]).resolve()  # 上传根目录
        self.allowed_extensions = {ext.lower() for ext in UPLOAD_CONFIG["allowed_extensions"]}  # 允许的文件扩展名集合
        self.max_file_size = int(UPLOAD_CONFIG["max_file_size"])  # 最大文件大小
        self.chunk_size = int(UPLOAD_CONFIG["chunk_size"])  # 分块大小
        self.chunk_overlap = int(UPLOAD_CONFIG["chunk_overlap"])  # 分块重叠大小
        self.pdf_multimodal_page_batch_size = int(UPLOAD_CONFIG.get("pdf_multimodal_page_batch_size", 2))
        self.pdf_multimodal_timeout_seconds = int(UPLOAD_CONFIG.get("pdf_multimodal_timeout_seconds", 90))  # 多模态超时时间
        self.image_use_multimodal_fallback = bool(
            UPLOAD_CONFIG.get("image_use_multimodal_fallback", True)
        )  # 图片 OCR 失败时是否走多模态
        self.image_force_multimodal_for_complex = bool(
            UPLOAD_CONFIG.get("image_force_multimodal_for_complex", True)
        )  # 复杂图片是否强制走多模态
        self.image_multimodal_min_ocr_chars = int(
            UPLOAD_CONFIG.get("image_multimodal_min_ocr_chars", 24)
        )  # OCR 结果过短阈值
        self.upload_root.mkdir(parents=True, exist_ok=True)  # 确保上传目录存在

    def upload_file(
        self,
        user_id: int,
        conversation_id: int,
        filename: str,
        content_type: str,
        file_bytes: bytes,
    ) -> UploadedFile:
        """
        上传并解析文件。
        
        流程：验证文件 -> 保存到磁盘 -> 解析内容 -> 分块 -> 保存到数据库 -> 同步到 Milvus。
        
        Args:
            user_id: 用户 ID
            conversation_id: 会话 ID
            filename: 文件名
            content_type: 文件 MIME 类型
            file_bytes: 文件字节数据
            
        Returns:
            UploadedFile: 上传文件记录
        """
        conversation = self._get_owned_conversation(user_id, conversation_id)  # 验证会话归属
        original_name = (filename or "").strip()
        if not original_name:
            raise ValueError("上传文件缺少文件名。")

        file_ext = Path(original_name).suffix.lower()  # 获取文件扩展名
        if file_ext not in self.allowed_extensions:
            raise ValueError(f"暂不支持该文件类型：{file_ext or '未知类型'}")

        if not file_bytes:
            raise ValueError("上传文件内容为空。")
        if len(file_bytes) > self.max_file_size:
            raise ValueError(f"文件过大，单个文件不能超过 {self.max_file_size // (1024 * 1024)} MB。")

        stored_path = self._build_storage_path(user_id, conversation_id, file_ext)  # 构建存储路径
        stored_path.parent.mkdir(parents=True, exist_ok=True)
        stored_path.write_bytes(file_bytes)  # 保存文件到磁盘

        # 创建上传文件记录
        uploaded_file = UploadedFile(
            user_id=user_id,
            conversation_id=conversation.id,
            original_name=original_name,
            stored_name=stored_path.name,
            file_ext=file_ext,
            mime_type=(content_type or "").strip() or None,
            size_bytes=len(file_bytes),
            storage_path=str(stored_path),
            parse_status="processing",  # 初始状态为处理中
            parse_error=None,
            text_length=0,
            chunk_count=0,
        )
        self.db.add(uploaded_file)
        self.db.commit()
        self.db.refresh(uploaded_file)

        try:
            parsed_text = self._parse_file(original_name, file_ext, file_bytes, stored_path=stored_path)  # 解析文件内容
            chunks = self._build_chunks(
                text=parsed_text,
                title=original_name,
                source=f"用户上传文件：{original_name}",
            )
            if not chunks:
                raise ValueError("文件解析后没有得到可检索文本，请检查文件内容是否为文本或可提取文字。")

            # 删除旧的分块记录
            self.db.query(UserDocumentChunk).filter(UserDocumentChunk.file_id == uploaded_file.id).delete(
                synchronize_session=False
            )
            # 保存新的分块记录
            for chunk in chunks:
                self.db.add(
                    UserDocumentChunk(
                        file_id=uploaded_file.id,
                        user_id=user_id,
                        conversation_id=conversation.id,
                        title=chunk["title"],
                        content=chunk["content"],
                        source=chunk["source"],
                        chunk_index=chunk["chunk_index"],
                    )
                )

            # 更新文件解析状态
            uploaded_file.parse_status = "ready"
            uploaded_file.parse_error = None
            uploaded_file.text_length = len(parsed_text)
            uploaded_file.chunk_count = len(chunks)
            self.db.commit()
            # 同步分块到 Milvus 向量库
            self._sync_chunks_to_milvus(
                file_id=uploaded_file.id,
                user_id=user_id,
                conversation_id=conversation.id,
                chunks=chunks,
            )
            self.db.refresh(uploaded_file)
            return uploaded_file
        except Exception as exc:
            uploaded_file.parse_status = "error"  # 解析失败
            uploaded_file.parse_error = str(exc)
            self.db.commit()
            raise

    def analyze_file_role(self, filename: str, file_bytes: bytes) -> Dict:
        """
        分析文件内容并推荐最适合的角色。
        
        Args:
            filename: 文件名
            file_bytes: 文件字节数据
            
        Returns:
            Dict: {role_type, confidence, reason, scores}
        """
        original_name = (filename or "").strip()
        if not original_name:
            raise ValueError("上传文件缺少文件名。")

        file_ext = Path(original_name).suffix.lower()
        if file_ext not in self.allowed_extensions:
            raise ValueError(f"暂不支持该文件类型：{file_ext or '未知类型'}")
        if not file_bytes:
            raise ValueError("上传文件内容为空。")
        if len(file_bytes) > self.max_file_size:
            raise ValueError(f"文件过大，单个文件不能超过 {self.max_file_size // (1024 * 1024)} MB。")

        parsed_text = self._parse_file(original_name, file_ext, file_bytes)  # 解析文件内容
        recommendation = self.recommend_role(parsed_text, original_name)  # 推荐角色
        return {
            **recommendation,
            "filename": original_name,
            "text_length": len(parsed_text),
        }

    def list_conversation_files(self, user_id: int, conversation_id: int) -> List[Dict]:
        """
        列出指定会话的所有上传文件。
        
        Args:
            user_id: 用户 ID
            conversation_id: 会话 ID
            
        Returns:
            List[Dict]: 文件信息列表
        """
        self._get_owned_conversation(user_id, conversation_id)
        files = (
            self.db.query(UploadedFile)
            .filter(
                UploadedFile.user_id == user_id,
                UploadedFile.conversation_id == conversation_id,
            )
            .order_by(UploadedFile.created_at.desc())  # 按创建时间降序排列
            .all()
        )
        return [self.serialize_file(item) for item in files]

    def delete_uploaded_file(self, user_id: int, file_id: int) -> None:
        """
        删除用户上传的文件。
        
        Args:
            user_id: 用户 ID
            file_id: 文件 ID
        """
        uploaded_file = (
            self.db.query(UploadedFile)
            .filter(UploadedFile.id == file_id, UploadedFile.user_id == user_id)
            .first()
        )
        if not uploaded_file:
            raise ValueError(f"文件不存在：{file_id}")

        self._delete_uploaded_file_record(uploaded_file)  # 删除文件记录和相关数据
        self.db.commit()

    def delete_conversation_files(self, conversation_id: int) -> None:
        """
        删除指定会话的所有文件。
        
        Args:
            conversation_id: 会话 ID
        """
        files = (
            self.db.query(UploadedFile)
            .filter(UploadedFile.conversation_id == conversation_id)
            .all()
        )
        for uploaded_file in files:
            self._delete_uploaded_file_record(uploaded_file)
        self.db.commit()

    def serialize_file(self, uploaded_file: UploadedFile) -> Dict:
        """
        将 UploadedFile 对象序列化为字典。
        
        Args:
            uploaded_file: UploadedFile 对象
            
        Returns:
            Dict: 文件信息字典
        """
        page_previews = self._list_file_page_previews(uploaded_file)
        return {
            "id": uploaded_file.id,
            "conversation_id": uploaded_file.conversation_id,
            "original_name": uploaded_file.original_name,
            "file_ext": uploaded_file.file_ext,
            "mime_type": uploaded_file.mime_type,
            "size_bytes": uploaded_file.size_bytes,
            "parse_status": uploaded_file.parse_status,
            "parse_error": uploaded_file.parse_error,
            "text_length": uploaded_file.text_length,
            "chunk_count": uploaded_file.chunk_count,
            "page_count": len(page_previews),
            "page_previews": page_previews,
            "created_at": uploaded_file.created_at,
            "updated_at": uploaded_file.updated_at,
        }

    def recommend_role(self, text: str, filename: str = "") -> Dict:
        """
        基于关键词匹配推荐最适合的角色。
        
        Args:
            text: 文件文本内容
            filename: 文件名（可选）
            
        Returns:
            Dict: {role_type, confidence, reason, scores}
        """
        corpus = f"{filename}\n{text or ''}".lower()  # 将文件名和文本合并为语料
        scores: Dict[str, int] = {}
        for role_type, keywords in self.ROLE_KEYWORDS.items():
            score = 0
            for keyword in keywords:
                score += corpus.count(keyword.lower())  # 统计每个关键词的出现次数
            scores[role_type] = score

        best_role = max(scores, key=scores.get)  # 得分最高的角色
        best_score = scores.get(best_role, 0)
        if best_score <= 0:  # 没有匹配到任何关键词
            return {
                "role_type": "custom_persona",
                "confidence": 0,
                "reason": "文件内容未匹配到固定角色关键词，建议使用全能型人格直接调用在线模型。",
                "scores": scores,
            }

        return {
            "role_type": best_role,
            "confidence": best_score,
            "reason": f"文件内容中与 {best_role} 相关的关键词匹配度最高。",
            "scores": scores,
        }

    def _get_owned_conversation(self, user_id: int, conversation_id: int) -> Conversation:
        """
        验证会话是否属于当前用户。
        
        Args:
            user_id: 用户 ID
            conversation_id: 会话 ID
            
        Returns:
            Conversation: 会话对象
        """
        conversation = (
            self.db.query(Conversation)
            .filter(Conversation.id == conversation_id, Conversation.user_id == user_id)
            .first()
        )
        if not conversation:
            raise ValueError("当前会话不存在，或不属于当前用户。")
        return conversation

    def _build_storage_path(self, user_id: int, conversation_id: int, file_ext: str) -> Path:
        """
        构建文件存储路径。
        
        Args:
            user_id: 用户 ID
            conversation_id: 会话 ID
            file_ext: 文件扩展名
            
        Returns:
            Path: 存储路径
        """
        folder = self.upload_root / f"user_{user_id}" / f"conversation_{conversation_id}"
        filename = f"{uuid.uuid4().hex}{file_ext}"  # 使用 UUID 生成唯一文件名
        return folder / filename

    def _delete_uploaded_file_record(self, uploaded_file: UploadedFile) -> None:
        """
        删除上传文件记录及其相关数据。
        
        清理流程：删除 Milvus 向量 -> 删除数据库分块记录 -> 删除磁盘文件 -> 删除数据库记录。
        
        Args:
            uploaded_file: UploadedFile 对象
        """
        self._delete_file_chunks_from_milvus(uploaded_file.id)  # 从 Milvus 删除向量
        self.db.query(UserDocumentChunk).filter(UserDocumentChunk.file_id == uploaded_file.id).delete(
            synchronize_session=False
        )
        file_path = Path(uploaded_file.storage_path or "")
        if file_path:
            try:
                resolved = file_path.resolve()
                if str(resolved).startswith(str(self.upload_root)):  # 确保路径在上传目录内
                    resolved.unlink(missing_ok=True)  # 删除磁盘文件
                    preview_dir = self._get_preview_dir(resolved)
                    if preview_dir.exists():
                        for item in preview_dir.iterdir():
                            if item.is_file():
                                item.unlink(missing_ok=True)
                        preview_dir.rmdir()
            except OSError:
                pass
        self.db.delete(uploaded_file)  # 删除数据库记录

    def _get_preview_dir(self, file_path: Path) -> Path:
        """返回上传文件对应的页图目录。"""
        return file_path.with_name(f"{file_path.stem}_pages")

    def _build_upload_static_url(self, file_path: Path) -> str:
        """将上传目录内的文件路径转换为静态访问 URL。"""
        try:
            relative = file_path.resolve().relative_to(self.upload_root.resolve())
        except Exception:
            return ""
        return "/uploads/" + "/".join(relative.parts)

    def _list_file_page_previews(self, uploaded_file: UploadedFile) -> List[Dict]:
        """返回页图预览列表（保留兼容）。"""
        if (uploaded_file.file_ext or "").lower() != ".pdf":
            return []

        file_path = Path(uploaded_file.storage_path or "")
        if not file_path:
            return []

        preview_dir = self._get_preview_dir(file_path)
        if not preview_dir.exists():
            return []

        previews: List[Dict] = []
        for image_path in sorted(preview_dir.glob("page_*.png")):
            match = re.search(r"page_(\d+)\.png$", image_path.name)
            page_number = int(match.group(1)) if match else 0
            previews.append(
                {
                    "page_number": page_number,
                    "image_url": self._build_upload_static_url(image_path),
                }
            )
        return previews

    def _sync_chunks_to_milvus(
        self,
        file_id: int,
        user_id: int,
        conversation_id: int,
        chunks: List[Dict],
    ) -> None:
        """
        将文件分块同步到 Milvus 向量库。
        
        MySQL 作为元数据存储，Milvus 作为向量检索引擎。
        
        Args:
            file_id: 文件 ID
            user_id: 用户 ID
            conversation_id: 会话 ID
            chunks: 分块列表
        """
        try:
            from vector_store import MilvusStore

            MilvusStore().insert_user_document_chunks(
                file_id=file_id,
                user_id=user_id,
                conversation_id=conversation_id,
                chunks=chunks,
            )
        except Exception:
            pass  # Milvus 不可用时静默失败

    def _delete_file_chunks_from_milvus(self, file_id: int) -> None:
        """
        从 Milvus 删除文件的分块向量。
        
        当文件被删除时，同步清理 Milvus 中的向量数据。
        
        Args:
            file_id: 文件 ID
        """
        try:
            from vector_store import MilvusStore

            MilvusStore().delete_user_file_chunks(file_id)
        except Exception:
            pass  # Milvus 不可用时静默失败

    def _parse_file(self, original_name: str, file_ext: str, file_bytes: bytes, stored_path: Path | None = None) -> str:
        """
        根据文件扩展名选择对应的解析方法。
        
        Args:
            original_name: 原始文件名
            file_ext: 文件扩展名
            file_bytes: 文件字节数据
            
        Returns:
            str: 解析后的文本内容
        """
        if file_ext in {".txt", ".md"}:  # 纯文本文件
            return self._parse_text_bytes(file_bytes)
        if file_ext == ".json":  # JSON 文件
            return self._parse_json(file_bytes)
        if file_ext == ".csv":  # CSV 文件
            return self._parse_csv(file_bytes, original_name)
        if file_ext == ".pdf":
            return self._parse_pdf(file_bytes, original_name, stored_path=stored_path)
        if file_ext == ".docx":  # DOCX 文件
            return self._parse_docx(file_bytes)
        if file_ext == ".xlsx":  # XLSX 文件
            return self._parse_xlsx(file_bytes, original_name)
        if file_ext in {".png", ".jpg", ".jpeg", ".bmp", ".webp"}:  # 图片文件
            return self._parse_image(file_bytes, original_name)
        raise ValueError(f"暂不支持该文件类型：{file_ext}")

    def _parse_text_bytes(self, file_bytes: bytes) -> str:
        """
        解析文本文件字节数据。
        
        尝试多种编码格式解码。
        
        Args:
            file_bytes: 文件字节数据
            
        Returns:
            str: 解析后的文本内容
        """
        text = ""
        for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):  # 尝试多种编码
            try:
                text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                text = ""
        cleaned = self.processor.clean_text(text)
        if not cleaned:
            raise ValueError("文本文件解析后内容为空。")
        return cleaned

    def _parse_json(self, file_bytes: bytes) -> str:
        """
        解析 JSON 文件。
        
        Args:
            file_bytes: 文件字节数据
            
        Returns:
            str: 格式化的 JSON 文本
        """
        text = self._parse_text_bytes(file_bytes)
        try:
            payload = json.loads(text)
            normalized = json.dumps(payload, ensure_ascii=False, indent=2)  # 格式化 JSON
        except json.JSONDecodeError:
            normalized = text
        cleaned = self.processor.clean_text(normalized)
        if not cleaned:
            raise ValueError("JSON 文件解析后内容为空。")
        return cleaned

    def _parse_csv(self, file_bytes: bytes, original_name: str) -> str:
        """
        解析 CSV 文件。
        
        Args:
            file_bytes: 文件字节数据
            original_name: 原始文件名
            
        Returns:
            str: 解析后的文本内容
        """
        decoded = self._parse_text_bytes(file_bytes)
        reader = csv.reader(io.StringIO(decoded))
        lines: List[str] = [f"文件：{original_name}"]
        for row_index, row in enumerate(reader, start=1):
            values = [str(value).strip() for value in row if str(value).strip()]
            if not values:
                continue
            lines.append(f"第 {row_index} 行：{' | '.join(values)}")
        cleaned = self.processor.clean_text("\n".join(lines))
        if not cleaned:
            raise ValueError("CSV 文件解析后内容为空。")
        return cleaned

    def _parse_pdf(self, file_bytes: bytes, original_name: str, stored_path: Path | None = None) -> str:
        """Prefer native PDF extraction and use multimodal parsing only for weak pages."""
        native_pages = self._extract_pdf_native_pages(file_bytes)
        rendered_pages = self._render_pdf_pages(file_bytes, stored_path=stored_path)
        if not rendered_pages:
            raise ValueError("PDF 没有生成可解析的页面。")

        rendered_by_page = {page["page_number"]: page for page in rendered_pages}
        page_count = len(rendered_pages)
        final_page_texts: Dict[int, str] = {}

        for native_page in native_pages:
            page_number = native_page["page_number"]
            if self._is_native_pdf_page_usable(native_page):
                final_page_texts[page_number] = native_page["content"]

        fallback_page_numbers = [
            page["page_number"]
            for page in rendered_pages
            if page["page_number"] not in final_page_texts
        ]

        if fallback_page_numbers:
            try:
                config = load_multimodal_llm_config()
                client = build_multimodal_openai_client(config)
            except Exception as exc:
                logger.warning("PDF 多模态兜底不可用，将使用原生抽取结果回退：%s", exc)
                client = None
                config = {}

            if client is not None:
                max_tokens = int(config.get("max_new_tokens") or config.get("max_tokens") or 3200)
                batch_size = max(self.pdf_multimodal_page_batch_size, 1)
                fallback_pages = [rendered_by_page[page_number] for page_number in fallback_page_numbers]

                for start in range(0, len(fallback_pages), batch_size):
                    batch = fallback_pages[start : start + batch_size]
                    batch_results = self._parse_pdf_batch_with_multimodal(
                        client=client,
                        config=config,
                        original_name=original_name,
                        pages=batch,
                        total_pages=page_count,
                        max_tokens=max_tokens,
                    )
                    for page in batch:
                        page_number = page["page_number"]
                        page_text = batch_results.get(page_number, "")
                        if page_text:
                            final_page_texts[page_number] = self._format_pdf_page_output(page_number, page_text)

        for native_page in native_pages:
            page_number = native_page["page_number"]
            if page_number not in final_page_texts and native_page["content"]:
                final_page_texts[page_number] = native_page["content"]

        ordered_text = [
            final_page_texts[page_number]
            for page_number in sorted(final_page_texts)
            if self.processor.clean_text(final_page_texts[page_number])
        ]
        merged = "\n\n".join(ordered_text).strip()
        if not merged:
            raise ValueError("PDF 解析后没有得到可用文本。")
        return merged

    def _extract_pdf_native_pages(self, file_bytes: bytes) -> List[Dict]:
        """Extract per-page text and tables directly from the PDF structure."""
        text_by_page = self._extract_pdf_text_blocks_native(file_bytes)
        tables_by_page = self._extract_pdf_tables_native(file_bytes)
        page_numbers = sorted(set(text_by_page) | set(tables_by_page))

        pages: List[Dict] = []
        for page_number in page_numbers:
            text_blocks = text_by_page.get(page_number, [])
            tables = tables_by_page.get(page_number, [])
            content = self._compose_native_pdf_page_content(page_number, text_blocks, tables)
            pages.append(
                {
                    "page_number": page_number,
                    "text_blocks": text_blocks,
                    "tables": tables,
                    "content": content,
                }
            )
        return pages

    def _extract_pdf_text_blocks_native(self, file_bytes: bytes) -> Dict[int, List[str]]:
        """Use PyMuPDF to extract ordered text blocks per page."""
        page_blocks: Dict[int, List[str]] = {}
        try:
            document = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as exc:
            raise ValueError(f"打开 PDF 失败：{exc}") from exc

        try:
            for index in range(document.page_count):
                page = document.load_page(index)
                blocks = page.get_text("blocks") or []
                ordered = sorted(blocks, key=lambda item: (round(item[1], 1), round(item[0], 1)))
                texts: List[str] = []
                for block in ordered:
                    block_text = self.processor.clean_text(block[4] or "")
                    if block_text:
                        texts.append(block_text)
                page_blocks[index + 1] = texts
        finally:
            document.close()

        return page_blocks

    def _extract_pdf_tables_native(self, file_bytes: bytes) -> Dict[int, List[str]]:
        """Use pdfplumber to extract tables and convert them into formal Markdown tables."""
        tables_by_page: Dict[int, List[str]] = {}
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                page_tables: List[str] = []
                try:
                    tables = page.extract_tables() or []
                except Exception:
                    tables = []

                for table in tables:
                    markdown = self._convert_pdf_table_to_markdown(table)
                    if markdown:
                        page_tables.append(markdown)
                tables_by_page[index] = page_tables
        return tables_by_page

    def _convert_pdf_table_to_markdown(self, table) -> str:
        """Normalize pdfplumber table output into a formal Markdown table."""
        if not table:
            return ""

        rows: List[List[str]] = []
        column_count = 0
        for row in table:
            normalized_row = [self.processor.clean_text(str(cell or "")) for cell in (row or [])]
            if not any(normalized_row):
                continue
            column_count = max(column_count, len(normalized_row))
            rows.append(normalized_row)

        if not rows or column_count < 2:
            return ""

        padded_rows: List[List[str]] = []
        for row in rows:
            padded = row + [""] * (column_count - len(row))
            padded_rows.append(padded)

        header = padded_rows[0]
        if not any(header):
            header = [f"列{index + 1}" for index in range(column_count)]

        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * column_count) + " |",
        ]
        for row in padded_rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)

    def _compose_native_pdf_page_content(self, page_number: int, text_blocks: List[str], tables: List[str]) -> str:
        """Compose a formal per-page output with body text and tables separated clearly."""
        sections: List[str] = [f"第 {page_number} 页"]

        body_blocks = [block for block in text_blocks if block]
        if body_blocks:
            sections.append("正文：")
            sections.extend(body_blocks)

        if tables:
            if body_blocks:
                sections.append("")
            sections.append("表格：")
            for table_index, markdown in enumerate(tables, start=1):
                sections.append(f"表 {table_index}：")
                sections.append(markdown)
                if table_index != len(tables):
                    sections.append("")

        return "\n".join(section for section in sections if section is not None).strip()

    def _is_native_pdf_page_usable(self, page_payload: Dict) -> bool:
        """Judge whether native extraction is good enough to avoid multimodal fallback."""
        content = self.processor.clean_text(page_payload.get("content", ""))
        tables = page_payload.get("tables") or []
        text_blocks = page_payload.get("text_blocks") or []
        if tables:
            return True

        if len(content) < 120:
            return False

        compact = re.sub(r"\s+", "", content)
        if len(compact) < 80:
            return False

        chinese_count = len(re.findall(r"[\u4e00-\u9fff]", compact))
        english_word_count = len(re.findall(r"[A-Za-z]{3,}", content))
        digit_count = len(re.findall(r"\d", compact))

        if chinese_count >= 20 or english_word_count >= 20:
            return True

        if digit_count > len(compact) * 0.65 and len(text_blocks) <= 3:
            return False

        return bool(text_blocks)

    def _format_pdf_page_output(self, page_number: int, page_text: str) -> str:
        """Wrap multimodal fallback output in the same formal page structure as native extraction."""
        cleaned = str(page_text or "").strip()
        if not cleaned:
            return ""

        if cleaned.startswith(f"第 {page_number} 页"):
            return cleaned

        return f"第 {page_number} 页\n正文：\n{cleaned}"

    def _render_pdf_pages(self, file_bytes: bytes, stored_path: Path | None = None) -> List[Dict]:
        """Render PDF pages and optionally persist preview PNG files."""
        preview_dir = self._get_preview_dir(stored_path) if stored_path else None
        if preview_dir:
            preview_dir.mkdir(parents=True, exist_ok=True)

        pages: List[Dict] = []
        try:
            document = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as exc:
            raise ValueError(f"打开 PDF 失败：{exc}") from exc

        try:
            for index in range(document.page_count):
                page = document.load_page(index)
                pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                preview_bytes = pix.tobytes("png")
                api_image_bytes = self._compress_pdf_page_for_api(preview_bytes)

                if preview_dir:
                    (preview_dir / f"page_{index + 1}.png").write_bytes(preview_bytes)

                pages.append(
                    {
                        "page_number": index + 1,
                        "preview_bytes": preview_bytes,
                        "api_image_bytes": api_image_bytes,
                    }
                )
        finally:
            document.close()

        return pages

    def _compress_pdf_page_for_api(self, image_bytes: bytes) -> bytes:
        """Resize/compress rendered PDF pages to reduce multimodal payload size."""
        with Image.open(io.BytesIO(image_bytes)) as image:
            image = image.convert("RGB")
            max_width = 1600
            if image.width > max_width:
                target_height = max(1, int(image.height * (max_width / float(image.width))))
                image = image.resize((max_width, target_height), Image.Resampling.LANCZOS)

            buffer = io.BytesIO()
            image.save(buffer, format="JPEG", quality=88, optimize=True)
            return buffer.getvalue()

    def _parse_pdf_batch_with_multimodal(
        self,
        client,
        config: Dict,
        original_name: str,
        pages: List[Dict],
        total_pages: int,
        max_tokens: int,
    ) -> Dict[int, str]:
        """Send a batch of rendered PDF pages to the multimodal API and split the response per page."""
        if not pages:
            return {}

        instruction = self._build_multimodal_pdf_instruction(
            original_name=original_name,
            page_numbers=[page["page_number"] for page in pages],
            total_pages=total_pages,
        )
        content: List[Dict] = [{"type": "text", "text": instruction}]
        for page in pages:
            image_base64 = base64.b64encode(page["api_image_bytes"]).decode("ascii")
            content.append({"type": "text", "text": f"Page marker: [[PAGE_{page['page_number']}]]"})
            content.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}})

        try:
            response = self._create_multimodal_completion(
                client=client,
                config=config,
                content=content,
                max_tokens=max_tokens,
                timeout=self.pdf_multimodal_timeout_seconds,
            )
        except Exception as exc:
            page_numbers = [page["page_number"] for page in pages]
            logger.warning("Multimodal PDF batch parsing failed for %s pages %s: %s", original_name, page_numbers, exc)
            return {}

        text = self._extract_multimodal_response_text(response)
        page_results = self._extract_pdf_page_results(text, [page["page_number"] for page in pages])
        cleaned_results: Dict[int, str] = {}
        for page_number, page_text in page_results.items():
            cleaned = self._clean_multimodal_markdown_text(page_text)
            cleaned = self._postprocess_pdf_page_text(cleaned)
            if not self._is_valid_pdf_multimodal_result(cleaned):
                logger.warning(
                    "Multimodal PDF parsing returned invalid text for %s page %s: %r",
                    original_name,
                    page_number,
                    cleaned[:200],
                )
                continue
            cleaned_results[page_number] = cleaned

        if not cleaned_results:
            logger.warning(
                "Multimodal PDF batch parsing returned no valid text for %s pages %s: %r",
                original_name,
                [page["page_number"] for page in pages],
                str(text)[:300],
            )
        return cleaned_results

    def _build_multimodal_pdf_instruction(self, original_name: str, page_numbers: List[int], total_pages: int) -> str:
        """Prompt for extracting cleaned text from one or more rendered PDF pages."""
        page_scope = ", ".join(str(number) for number in page_numbers)
        return (
            f"You are extracting text from pages {page_scope} of {total_pages} in the PDF file "
            f"'{original_name}'.\n"
            "Return the result page by page using this exact wrapper format:\n"
            "[[PAGE_1]] ... [[/PAGE_1]]\n"
            "Replace 1 with the real page number for each provided page.\n"
            "Strict rules:\n"
            "1. Remove headers, footers, page numbers, repeated watermarks, scan stamps, and obvious garbled noise.\n"
            "2. Preserve the original reading order of the main body text.\n"
            "3. Keep heading hierarchy stable. Preserve section titles, subsection titles, numbered headings, and list levels when visible.\n"
            "4. For tables, output them as Markdown tables with pipe separators. Always keep one row per line.\n"
            "5. In tables, keep column order, row labels, units, percentages, negative signs, dates, and totals exactly as shown. Never merge adjacent cells or adjacent rows.\n"
            "6. If the table header contains years, periods, or forecast labels such as 2023, 2024, 2025E, H1, Q1, or FY, place each one in its own column cell.\n"
            "7. For financial reports, preserve line items, period labels, currencies, share counts, notes references, and subtotal/total relationships. Do not merge adjacent rows. Do not concatenate neighboring numbers.\n"
            "8. If a number is visually separated into different columns, keep those numbers in different columns in the Markdown table.\n"
            "9. For rows like revenue, net profit, EPS, PE, PB, EV/EBITDA, ROE, margin, or growth rate, keep the metric name in the first column and each period value in later columns.\n"
            "10. For academic papers, preserve title, authors if visible on the page, section headings, figure/table captions, equations, footnotes, references, and citation markers.\n"
            "11. Footnotes and end-of-page notes should be kept only if they are readable and meaningful, and they should appear after the main body of that page under a 'Footnotes:' label.\n"
            "12. If a line is clearly garbled or meaningless, drop it instead of guessing.\n"
            "13. Do not summarize, explain, translate, or add information that is not visible on the page.\n"
            "14. Return plain text only. Do not use Markdown code fences.\n"
            "15. For images, diagrams, flowcharts, and charts, keep figure titles/captions if present; otherwise mark as [图片] or [无法识别图片].\n"
            "16. Preserve ordered lists (1., 2.) and unordered lists (-, •) with original symbols and indentation; list items stay on separate lines.\n"
            "17. If a page contains no meaningful content, output [[PAGE_X]] [空白页] [[/PAGE_X]].\n"
        )

    def _extract_pdf_page_results(self, response_text: str, page_numbers: List[int]) -> Dict[int, str]:
        """Split a batched multimodal PDF response back into per-page text blocks."""
        results: Dict[int, str] = {}
        raw_text = str(response_text or "").replace("\r", "\n")
        for page_number in page_numbers:
            pattern = rf"\[\[PAGE_{page_number}\]\](.*?)\[\[/PAGE_{page_number}\]\]"
            match = re.search(pattern, raw_text, flags=re.DOTALL | re.IGNORECASE)
            if match:
                results[page_number] = match.group(1).strip()
        if results:
            return results

        # Fallback for providers that ignore the closing marker format but keep page headers.
        for index, page_number in enumerate(page_numbers):
            start_pattern = rf"\[\[PAGE_{page_number}\]\]"
            start_match = re.search(start_pattern, raw_text, flags=re.IGNORECASE)
            if not start_match:
                continue

            start = start_match.end()
            end = len(raw_text)
            for next_page_number in page_numbers[index + 1 :]:
                next_match = re.search(rf"\[\[PAGE_{next_page_number}\]\]", raw_text[start:], flags=re.IGNORECASE)
                if next_match:
                    end = start + next_match.start()
                    break
            results[page_number] = raw_text[start:end].strip()

        return results

    def _is_valid_pdf_multimodal_result(self, text: str) -> bool:
        """PDF page results may be English or Chinese, so validate by content length only."""
        cleaned = self.processor.clean_text(text or "")
        if not cleaned:
            return False

        compact = re.sub(r"\s+", "", cleaned).lower()
        if not compact:
            return False

        for token in self.INVALID_MULTIMODAL_TOKENS:
            compact = compact.replace(token, "")

        if len(compact) < 24:
            return False

        alnum_chars = re.findall(r"[a-z0-9\u4e00-\u9fff]", compact)
        return len(alnum_chars) >= 12

    def _postprocess_pdf_page_text(self, text: str) -> str:
        """Repair compact financial table text into Markdown tables when possible."""
        lines = [line.rstrip() for line in str(text or "").splitlines()]
        if not lines:
            return ""

        rewritten: List[str] = []
        index = 0
        while index < len(lines):
            line = lines[index].strip()
            header_tokens = self._extract_compact_year_header_tokens(line)
            if header_tokens:
                table_rows: List[List[str]] = []
                lookahead = index + 1
                while lookahead < len(lines):
                    parsed_row = self._parse_compact_financial_row(lines[lookahead], len(header_tokens))
                    if not parsed_row:
                        break
                    table_rows.append(parsed_row)
                    lookahead += 1

                if len(table_rows) >= 2:
                    rewritten.extend(self._build_markdown_table_lines(header_tokens, table_rows))
                    index = lookahead
                    continue

            rewritten.append(line)
            index += 1

        return "\n".join(line for line in rewritten if line is not None).strip()

    def _extract_compact_year_header_tokens(self, line: str) -> List[str]:
        """Detect dense year/period header rows such as 202320242025E2026E2027E."""
        compact = re.sub(r"[\s|]+", "", str(line or ""))
        if not compact:
            return []

        tokens = re.findall(r"(?:20\d{2}E?|20\d{2}A?|20\d{2}|FY\d{2,4}E?|H[12]\d{2,4}E?|Q[1-4]\d{2,4}E?)", compact, flags=re.IGNORECASE)
        if len(tokens) < 3:
            return []

        return tokens if "".join(tokens).lower() == compact.lower() else []

    def _parse_compact_financial_row(self, line: str, expected_value_count: int) -> List[str] | None:
        """Split compact financial metric rows into label + per-period values."""
        raw = str(line or "").strip()
        if not raw or "|" in raw:
            return None

        compact = re.sub(r"\s+", "", raw)
        match = re.match(r"^([^0-9+\-−]+?)([+\-−]?\d.*)$", compact)
        if not match:
            return None

        label = match.group(1).strip()
        value_text = match.group(2).strip()
        if not label or not value_text:
            return None

        values = self._extract_compact_metric_values(value_text, expected_value_count)
        if len(values) != expected_value_count:
            return None

        return [label, *values]

    def _extract_compact_metric_values(self, value_text: str, expected_count: int) -> List[str]:
        """Best-effort split for dense numeric sequences extracted from financial tables."""
        compact = re.sub(r"\s+", "", str(value_text or ""))
        if not compact or expected_count <= 0:
            return []

        direct_pattern = r"[+\-−]?\d{1,3}(?:,\d{3})+(?:\.\d{1,2})?%?|[+\-−]?\d+(?:\.\d{1,2})?%?"
        direct_matches = re.findall(direct_pattern, compact)
        if direct_matches and "".join(direct_matches) == compact and len(direct_matches) == expected_count:
            return direct_matches

        if compact.isdigit() and len(compact) % expected_count == 0:
            chunk_len = len(compact) // expected_count
            if 2 <= chunk_len <= 6:
                return [compact[i : i + chunk_len] for i in range(0, len(compact), chunk_len)]

        memo: Dict[tuple[str, int], List[str] | None] = {}

        def solve(remaining: str, count: int) -> List[str] | None:
            key = (remaining, count)
            if key in memo:
                return memo[key]

            if count == 0:
                memo[key] = [] if not remaining else None
                return memo[key]
            if not remaining:
                memo[key] = None
                return None

            if count == 1:
                if re.fullmatch(direct_pattern, remaining):
                    memo[key] = [remaining]
                    return memo[key]
                memo[key] = None
                return None

            for candidate in self._iter_compact_numeric_candidates(remaining):
                tail = solve(remaining[len(candidate):], count - 1)
                if tail is not None:
                    memo[key] = [candidate, *tail]
                    return memo[key]

            memo[key] = None
            return None

        return solve(compact, expected_count) or []

    def _iter_compact_numeric_candidates(self, text: str) -> List[str]:
        """Enumerate plausible numeric tokens from the start of a dense sequence."""
        if not text:
            return []

        normalized = text.replace("−", "-")
        candidates: List[str] = []

        comma_match = re.match(r"[+\-]?\d{1,3}(?:,\d{3})+(?:\.\d{1,2})?%?", normalized)
        if comma_match:
            candidates.append(comma_match.group(0))
            return candidates

        integer_match = re.match(r"[+\-]?\d+", normalized)
        if not integer_match:
            return []

        integer_token = integer_match.group(0)
        next_index = len(integer_token)

        if next_index < len(normalized) and normalized[next_index] == ".":
            for decimal_len in (1, 2):
                end = next_index + 1 + decimal_len
                if end <= len(normalized) and normalized[next_index + 1 : end].isdigit():
                    token = normalized[:end]
                    candidates.append(token)
                    if end < len(normalized) and normalized[end] == "%":
                        candidates.append(normalized[: end + 1])

        if next_index < len(normalized) and normalized[next_index] == "%":
            candidates.append(normalized[: next_index + 1])

        if len(integer_token.lstrip("+-")) <= 6:
            candidates.append(integer_token)

        deduped: List[str] = []
        seen = set()
        for candidate in candidates:
            if candidate not in seen:
                deduped.append(candidate)
                seen.add(candidate)
        return deduped

    def _build_markdown_table_lines(self, header_tokens: List[str], rows: List[List[str]]) -> List[str]:
        """Build Markdown table lines from parsed compact financial rows."""
        table_lines = [
            "| 指标 | " + " | ".join(header_tokens) + " |",
            "| --- | " + " | ".join(["---"] * len(header_tokens)) + " |",
        ]
        for row in rows:
            table_lines.append("| " + " | ".join(row) + " |")
        return table_lines

    def _create_multimodal_completion(self, client, config: Dict, content: List[Dict], max_tokens: int, timeout: float):
        """统一创建多模态请求，优先走 responses 接口，不支持时回退 chat.completions。"""
        response_input = [
            {
                "role": "user",
                "content": [self._to_multimodal_response_input_item(item) for item in content],
            }
        ]

        response_kwargs = {
            "model": config["model_name"],
            "input": response_input,
            "max_output_tokens": max_tokens,
            "temperature": float(config.get("temperature", 0.05)),
            "top_p": float(config.get("top_p", 0.7)),
            "timeout": timeout,
        }

        repetition_penalty = config.get("repetition_penalty")

        try:
            responses_api = getattr(client, "responses", None)
            if responses_api is not None:
                return responses_api.create(**response_kwargs)
        except Exception as exc:
            logger.warning("Multimodal responses API failed, falling back to chat.completions: %s", exc)

        request_kwargs = {
            "model": config["model_name"],
            "messages": [{"role": "user", "content": content}],
            "max_tokens": max_tokens,
            "temperature": float(config.get("temperature", 0.05)),
            "top_p": float(config.get("top_p", 0.7)),
            "timeout": timeout,
        }

        if repetition_penalty is not None:
            request_kwargs["extra_body"] = {
                "repetition_penalty": float(repetition_penalty),
            }

        try:
            return client.chat.completions.create(**request_kwargs)
        except Exception as exc:
            if "extra_body" in request_kwargs:
                logger.warning(
                    "Multimodal request rejected repetition_penalty, retrying without it: %s",
                    exc,
                )
                request_kwargs.pop("extra_body", None)
                return client.chat.completions.create(**request_kwargs)
            raise

    def _to_multimodal_response_input_item(self, item: Dict) -> Dict:
        """Map chat-style multimodal items into responses-style input items."""
        item_type = str(item.get("type") or "").strip()
        if item_type == "text":
            return {
                "type": "input_text",
                "text": str(item.get("text") or ""),
            }

        if item_type == "image_url":
            image_url = item.get("image_url")
            if isinstance(image_url, dict):
                url = str(image_url.get("url") or "")
            else:
                url = str(image_url or "")
            return {
                "type": "input_image",
                "image_url": url,
            }

        return item

    def _clean_multimodal_markdown_text(self, text: str) -> str:
        """对多模态结构化结果做温和清洗，尽量保留 Markdown 和表格结构。"""
        if not text:
            return ""

        cleaned = str(text).replace("\r", "\n")
        cleaned = re.sub(r"[ \t\f\v]+", " ", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

        for token in self.INVALID_MULTIMODAL_TOKENS:
            cleaned = cleaned.replace(token, "")

        cleaned = re.sub(
            r"[^\u4e00-\u9fa5a-zA-Z0-9，。；：！？、（）《》“”‘’【】\[\]\-—/#\\\n\.%,:;!?+|*_`~=<>()]",
            "",
            cleaned,
        )
        return cleaned.strip()

    def _extract_multimodal_response_text(self, response) -> str:
        """
        从多模态 LLM 响应中提取文本。
        
        Args:
            response: LLM 响应对象
            
        Returns:
            str: 提取的文本
        """
        texts: List[str] = []

        output = getattr(response, "output", None) or []
        for item in output:
            for content in getattr(item, "content", None) or []:
                text_value = getattr(content, "text", None)
                if text_value:
                    texts.append(str(text_value))

        output_text = getattr(response, "output_text", None)
        if output_text:
            texts.append(str(output_text))

        response_output = getattr(response, "output", None) or []
        for item in response_output:
            item_type = getattr(item, "type", None)
            if item_type == "message":
                for content_item in getattr(item, "content", None) or []:
                    content_type = getattr(content_item, "type", None)
                    if content_type in {"output_text", "text"}:
                        text_value = getattr(content_item, "text", None)
                        if text_value:
                            texts.append(str(text_value))
                    elif isinstance(content_item, dict):
                        text_value = content_item.get("text")
                        if text_value:
                            texts.append(str(text_value))
            elif isinstance(item, dict):
                for content_item in item.get("content") or []:
                    if isinstance(content_item, dict):
                        text_value = content_item.get("text")
                        if text_value:
                            texts.append(str(text_value))

        choices = getattr(response, "choices", None) or []
        for choice in choices:
            message = getattr(choice, "message", None)
            if not message:
                continue
            content = getattr(message, "content", None)
            if isinstance(content, str) and content.strip():
                texts.append(content)
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, dict):
                        text_value = item.get("text")
                        if text_value:
                            texts.append(str(text_value))

        return "\n".join(texts).strip()

    def _is_valid_multimodal_result(self, text: str) -> bool:
        """
        验证多模态解析结果是否有效。
        
        Args:
            text: 待验证的文本
            
        Returns:
            bool: 是否有效
        """
        cleaned = self.processor.clean_text(text or "")
        if not cleaned:
            return False

        compact = re.sub(r"\s+", "", cleaned).lower()
        if not compact:
            return False

        for token in self.INVALID_MULTIMODAL_TOKENS:
            compact = compact.replace(token, "")

        if len(compact) < 24:  # 文本太短视为无效
            return False

        chinese_chars = re.findall(r"[\u4e00-\u9fff]", compact)
        if len(chinese_chars) < 6:  # 中文字符太少视为无效
            return False

        return True

    def _parse_docx(self, file_bytes: bytes) -> str:
        """
        解析 DOCX 文件。
        
        Args:
            file_bytes: 文件字节数据
            
        Returns:
            str: 解析后的文本内容
        """
        document = docx.Document(io.BytesIO(file_bytes))
        parts: List[str] = []

        for paragraph in document.paragraphs:  # 提取段落文本
            text = self.processor.clean_text(paragraph.text or "")
            if text:
                parts.append(text)

        for table_index, table in enumerate(document.tables, start=1):  # 提取表格文本
            row_texts: List[str] = []
            for row in table.rows:
                cells = [self.processor.clean_text(cell.text or "") for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    row_texts.append(" | ".join(cells))
            if row_texts:
                parts.append(f"表格 {table_index}\n" + "\n".join(row_texts))

        if not parts:
            raise ValueError("DOCX 文件解析后内容为空。")
        return "\n\n".join(parts)

    def _parse_xlsx(self, file_bytes: bytes, original_name: str) -> str:
        """
        解析 XLSX 文件。
        
        Args:
            file_bytes: 文件字节数据
            original_name: 原始文件名
            
        Returns:
            str: 解析后的文本内容
        """
        workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        parts: List[str] = [f"文件：{original_name}"]
        for sheet in workbook.worksheets:  # 遍历所有工作表
            rows: List[str] = []
            for row in sheet.iter_rows(values_only=True):
                values = [self.processor.clean_text(str(cell or "")) for cell in row]
                values = [value for value in values if value]
                if values:
                    rows.append(" | ".join(values))
            if rows:
                parts.append(f"工作表：{sheet.title}\n" + "\n".join(rows))

        cleaned = self.processor.clean_text("\n\n".join(parts))
        if not cleaned:
            raise ValueError("XLSX 文件解析后内容为空。")
        return cleaned

    def _build_chunks(self, text: str, title: str, source: str) -> List[Dict]:
        """
        将文本分割成重叠的分块。
        
        Args:
            text: 待分块文本
            title: 标题
            source: 来源
            
        Returns:
            List[Dict]: 分块列表，每块包含 title、content、source、chunk_index
        """
        cleaned = self.processor.clean_text(text)
        if not cleaned:
            return []

        units = self._split_document_units(cleaned)  # 按文档结构分割
        if not units:
            units = [cleaned]

        chunks: List[str] = []
        current = ""
        for unit in units:
            candidate = f"{current}{unit}" if not current else f"{current}\n{unit}"
            if current and len(candidate) > self.chunk_size:  # 超过块大小限制时切分
                chunks.append(current)
                overlap_text = current[-self.chunk_overlap :] if self.chunk_overlap else ""
                current = f"{overlap_text}\n{unit}".strip()
            else:
                current = candidate
        if current:
            chunks.append(current)

        result: List[Dict] = []
        for index, chunk in enumerate(chunks):
            normalized_chunk = self.processor.clean_text(chunk)
            if not normalized_chunk:
                continue
            result.append(
                {
                    "title": title,
                    "content": normalized_chunk,
                    "source": source,
                    "chunk_index": index,
                }
            )
        return result

    def _split_document_units(self, text: str) -> List[str]:
        """
        按文档结构（页、正文、表格、OCR 标记）分割文本单元。
        
        Args:
            text: 待分割文本
            
        Returns:
            List[str]: 文本单元列表
        """
        units: List[str] = []
        current: List[str] = []

        for raw_line in text.splitlines():
            line = self.processor.clean_text(raw_line)
            if not line:
                if current:
                    units.append("\n".join(current).strip())
                    current = []
                continue

            # 检测到文档结构标记（页、正文、表格、OCR）时切分
            if line.startswith("[第") or line.startswith("[正文]") or line.startswith("[表格") or line.startswith("[OCR]"):
                if current:
                    units.append("\n".join(current).strip())
                    current = []
            current.append(line)

        if current:
            units.append("\n".join(current).strip())

        return units

    def _parse_image(self, file_bytes: bytes, original_name: str) -> str:
        """
        解析图片文件。
        
        策略：
        - 普通文字图片优先走 OCR。
        - 图表/截图/低文字复杂图片优先走在线多模态解析。
        - OCR 失败或结果过短时，可回退到多模态解析。
        
        Args:
            file_bytes: 图片字节数据
            original_name: 原始文件名
            
        Returns:
            str: 解析后的文本内容
        """
        complexity_summary = self._inspect_image_complexity(file_bytes)
        parsed_text = self._extract_text_from_image_bytes(
            file_bytes,
            original_name=original_name,
            complexity_summary=complexity_summary,
        )
        if parsed_text:
            return parsed_text
        raise ValueError("图片解析失败：OCR 与在线多模态均未返回有效内容。")

    def _inspect_image_complexity(self, file_bytes: bytes) -> Dict:
        """
        检测图片是否属于复杂视觉内容。

        Args:
            file_bytes: 图片字节数据

        Returns:
            Dict: 复杂度摘要
        """
        summary = {
            "force_multimodal": False,
            "reason": "未命中复杂图片条件",
            "width": 0,
            "height": 0,
        }
        if not self.image_force_multimodal_for_complex:
            return summary

        try:
            with Image.open(io.BytesIO(file_bytes)) as image:
                width, height = image.size
        except Exception:
            return summary

        summary["width"] = width
        summary["height"] = height

        # 超宽、超高、近似长截图时优先走多模态
        if width >= 1800 or height >= 2200:
            summary["force_multimodal"] = True
            summary["reason"] = f"图片分辨率较大（{width}x{height}）"
        elif height >= width * 1.8:
            summary["force_multimodal"] = True
            summary["reason"] = f"图片疑似长截图（{width}x{height}）"

        return summary

    def _extract_text_from_image_bytes(
        self,
        file_bytes: bytes,
        original_name: str,
        complexity_summary: Dict | None = None,
    ) -> str:
        """
        统一的图片解析策略：
        - 普通文字图先走 OCR
        - 复杂图片/长截图/图表先走在线多模态
        - OCR 失败或 OCR 结果太短时回退到在线多模态
        """
        summary = complexity_summary or self._inspect_image_complexity(file_bytes)

        if summary.get("force_multimodal"):
            return self._parse_image_with_multimodal(
                file_bytes,
                original_name=original_name,
                complexity_summary=summary,
            )

        try:
            lines = self._run_ocr_worker(file_bytes)
        except Exception as exc:
            logger.warning("OCR failed for %s: %s", original_name, exc)
            if self.image_use_multimodal_fallback:
                return self._parse_image_with_multimodal(
                    file_bytes,
                    original_name=original_name,
                    complexity_summary={**summary, "reason": f"OCR 失败：{exc}"},
                )
            return ""

        if not lines:
            if self.image_use_multimodal_fallback:
                return self._parse_image_with_multimodal(
                    file_bytes,
                    original_name=original_name,
                    complexity_summary={**summary, "reason": "OCR 未识别到有效文字"},
                )
            return ""

        ocr_text = "\n".join(lines).strip()
        compact_ocr_text = re.sub(r"\s+", "", ocr_text)
        if len(compact_ocr_text) < self.image_multimodal_min_ocr_chars and self.image_use_multimodal_fallback:
            multimodal_text = self._parse_image_with_multimodal(
                file_bytes,
                original_name=original_name,
                complexity_summary={**summary, "reason": f"OCR 结果过短（{len(compact_ocr_text)} 字）"},
            )
            if multimodal_text:
                return multimodal_text

        return ocr_text

    def _parse_image_with_multimodal(self, file_bytes: bytes, original_name: str, complexity_summary=None) -> str:
        """
        使用在线多模态模型解析图片。

        Args:
            file_bytes: 图片字节数据
            original_name: 原始文件名
            complexity_summary: 复杂度摘要

        Returns:
            str: 解析后的文本内容，失败返回空字符串
        """
        try:
            config = load_multimodal_llm_config()
            client = build_multimodal_openai_client(config)
        except Exception as exc:
            logger.warning("Multimodal image client unavailable: %s", exc)
            return ""

        image_base64 = base64.b64encode(file_bytes).decode("ascii")
        instruction = self._build_multimodal_image_instruction(original_name, complexity_summary)
        content: List[Dict] = [
            {"type": "text", "text": instruction},
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_base64}"}},
        ]

        try:
            response = self._create_multimodal_completion(
                client=client,
                config=config,
                content=content,
                max_tokens=int(config.get("max_new_tokens") or config.get("max_tokens") or 2400),
                timeout=self.pdf_multimodal_timeout_seconds,
            )
        except Exception as exc:
            logger.warning("Multimodal image parsing request failed: %s", exc)
            return ""

        text = self._extract_multimodal_response_text(response)
        cleaned = self._clean_multimodal_markdown_text(text)
        if not self._is_valid_multimodal_result(cleaned):
            logger.warning("Multimodal image parsing returned invalid text: %r", cleaned[:200])
            return ""
        return cleaned

    def _build_multimodal_image_instruction(self, original_name: str, complexity_summary=None) -> str:
        """
        构建图片多模态解析提示词。

        Args:
            original_name: 原始文件名
            complexity_summary: 复杂度摘要

        Returns:
            str: 提示词
        """
        complexity_hint = ""
        if complexity_summary and complexity_summary.get("reason"):
            complexity_hint = f"这张图片的处理重点：{complexity_summary['reason']}。"

        return (
            f"你现在是中文图片内容提取助手，正在处理图片文件《{original_name}》。"
            f"{complexity_hint}"
            "请严格依据图片可见内容输出，不要编造图片外信息。"
            "如果图片里主要是文字，请按原文顺序转写；如果是表格，请按行转写字段、数值和单位；"
            "如果是图表，只允许转写图标题、坐标轴名称、图例、标注数值和图片中已经写出的说明文字。"
            "不要自行概括趋势，不要补充结论，不要写无关客套话。"
            "输出必须是中文纯文本，不要写 Markdown 代码块。"
        )

    def _run_ocr_worker(self, image_bytes: bytes) -> List[str]:
        """
        运行 OCR 工作进程解析图片。
        
        使用子进程运行 ocr_worker.py，避免 protobuf 与 Milvus 的冲突。
        
        Args:
            image_bytes: 图片字节数据
            
        Returns:
            List[str]: OCR 识别结果行列表
        """
        worker_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ocr_worker.py")
        if not os.path.exists(worker_path):
            raise RuntimeError(f"OCR worker 脚本不存在：{worker_path}")

        completed = subprocess.run(
            [sys.executable, worker_path],
            input=image_bytes,
            capture_output=True,
            check=False,
            timeout=120,
        )
        if completed.returncode != 0:
            stderr = (completed.stderr or b"").decode("utf-8", errors="replace").strip()
            stdout = (completed.stdout or b"").decode("utf-8", errors="replace").strip()
            raise RuntimeError(stderr or stdout or f"OCR worker 退出码：{completed.returncode}")

        output = (completed.stdout or b"").decode("utf-8", errors="replace").strip()
        if not output:
            return []

        try:
            payload = json.loads(output)
        except json.JSONDecodeError:
            # 兼容旧格式：worker 直接输出逐行文本
            return [line.strip() for line in output.splitlines() if line.strip()]

        if isinstance(payload, dict):
            lines = payload.get("lines") or []
            if isinstance(lines, list):
                return [self.processor.clean_text(str(line or "")) for line in lines if self.processor.clean_text(str(line or ""))]

        raise RuntimeError(f"OCR worker 返回了无法识别的结果格式：{output[:200]}")
